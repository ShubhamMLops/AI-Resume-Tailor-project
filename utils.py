from __future__ import annotations
from typing import List, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, ListFlowable, ListItem, KeepTogether

HEADING_COLOR_HEX = "#1F4E79"  # deep blue for headings
HEADING_COLOR_RGB = (0x1F, 0x4E, 0x79)

KNOWN_SECTIONS = {
    "Profile Summary","Core Skills","Core Competencies","Technical Skills",
    "Work Experience","Education","Certifications","Projects"
}

def export_docx(text: str, out_path: str) -> str:
    doc = Document()
    styles = doc.styles
    # Custom colored heading style
    if "TailorHeading" not in [s.name for s in styles]:
        hstyle = styles.add_style("TailorHeading", WD_STYLE_TYPE.PARAGRAPH)
        hstyle.font.name = "Calibri"
        hstyle.font.size = Pt(14)
        hstyle.font.bold = True
        hstyle.font.color.rgb = RGBColor(*HEADING_COLOR_RGB)
    else:
        hstyle = styles["TailorHeading"]
    nstyle = styles["Normal"]
    nstyle.font.name = "Calibri"
    nstyle.font.size = Pt(11)

    for raw in text.splitlines():
        s = raw.rstrip()
        if not s:
            doc.add_paragraph("")
            continue
        if s in KNOWN_SECTIONS:
            p = doc.add_paragraph()
            p.style = hstyle
            p.add_run(s)
        elif s.startswith("• "):
            p = doc.add_paragraph(s[2:])
            p.style = doc.styles["List Bullet"]
        elif s.startswith("- "):
            p = doc.add_paragraph(s[2:])
            p.style = doc.styles["List Bullet"]
        else:
            doc.add_paragraph(s)
    doc.save(out_path)
    return out_path

def _mk_styles():
    styles = getSampleStyleSheet()
    if "HeadingX" not in styles:
        styles.add(ParagraphStyle(
            name="HeadingX", parent=styles["Heading2"],
            fontName="Helvetica-Bold", fontSize=12.5,
            spaceBefore=8, spaceAfter=6,
            textColor=colors.HexColor(HEADING_COLOR_HEX)
        ))
    if "BodyX" not in styles:
        styles.add(ParagraphStyle(
            name="BodyX", parent=styles["Normal"],
            fontName="Helvetica", fontSize=10.5, leading=13
        ))
    return styles

def _parse_lines(text: str) -> List[Tuple[str,str]]:
    out = []
    for raw in text.splitlines():
        s = raw.rstrip()
        if not s:
            out.append(("blank",""))
        elif s in KNOWN_SECTIONS:
            out.append(("heading", s))
        elif s.startswith("• ") or s.startswith("- "):
            out.append(("bullet", s[2:]))
        elif s.startswith("---"):
            out.append(("hr",""))
        else:
            out.append(("para", s))
    return out

def export_pdf(text: str, out_path: str) -> str:
    doc = SimpleDocTemplate(out_path, pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=18*mm, bottomMargin=18*mm)
    styles = _mk_styles()
    story = []
    bullets = []
    def flush_bullets():
        nonlocal bullets, story
        if bullets:
            lst = ListFlowable([ListItem(Paragraph(b, styles["BodyX"])) for b in bullets],
                               bulletType="bullet", leftIndent=12,
                               bulletFontName="Helvetica", bulletFontSize=9)
            story.append(KeepTogether(lst))
            story.append(Spacer(1, 4))
            bullets = []
    for t, txt in _parse_lines(text):
        if t == "heading":
            flush_bullets()
            story.append(Spacer(1, 4))
            story.append(Paragraph(txt, styles["HeadingX"]))
            story.append(HRFlowable(width="100%", thickness=0.6, lineCap="round", color=colors.HexColor("#DDDDDD")))
            story.append(Spacer(1, 6))
        elif t == "bullet":
            bullets.append(txt)
        elif t == "hr":
            flush_bullets()
            story.append(HRFlowable(width="100%", thickness=0.6, lineCap="round", color=colors.HexColor("#DDDDDD")))
            story.append(Spacer(1, 6))
        elif t == "para":
            flush_bullets()
            if txt:
                story.append(Paragraph(txt, styles["BodyX"]))
        else:  # blank
            flush_bullets()
            story.append(Spacer(1, 4))
    flush_bullets()
    doc.build(story)
    return out_path
