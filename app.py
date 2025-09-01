import io, os, json, re
import streamlit as st
from utils import export_docx, export_pdf
from pipeline import extract_gaps
from pipeline import (
    analyze,
    tailor,
    extract_keywords_llm,
    extract_contacts_llm,
    polish_keyword_sentences,
    sanitize_markdown,
    extract_ats_llm_from_optimizer,
    generate_keyword_sentences,
    generate_summary_bullets,
    bulletize_summary_preserve_meaning,   # <-- add this
)


# -------------------------
# Stable editor/download state
# -------------------------
if "tailored_text" not in st.session_state:
    st.session_state["tailored_text"] = ""
if "tailored_edit" not in st.session_state:
    st.session_state["tailored_edit"] = ""
if "tailored_saved" not in st.session_state:
    st.session_state["tailored_saved"] = False
if "kw_sentences_edit" not in st.session_state:
    st.session_state["kw_sentences_edit"] = ""
if "kw_sentences_saved_text" not in st.session_state:
    st.session_state["kw_sentences_saved_text"] = ""

# -------------------------
# Token helpers
# -------------------------
_TOKEN_RE = re.compile(r"[A-Za-z0-9#+.]+")
def _tok_seq(s: str):
    return [t.lower() for t in _TOKEN_RE.findall(s or "")]

def _canon(s: str) -> str:
    return " ".join(_tok_seq(s))

def _present_line(base_text: str, line: str) -> bool:
    """
    True if 'line' already exists in the resume:
      - exact token sequence
      - compacted match (e.g., CI/CD -> cicd)
      - simple plural/singular for single word >3 chars
    """
    base_tokens = _tok_seq(base_text)
    base_compact = "".join(base_tokens)
    kt = _tok_seq(line)
    if not kt:
        return True
    L = len(kt)
    for i in range(0, len(base_tokens) - L + 1):
        if base_tokens[i:i+L] == kt:
            return True
    if "".join(kt) in base_compact:
        return True
    if L == 1 and len(kt[0]) > 3:
        w = kt[0]
        alt = w[:-1] if w.endswith("s") else w + "s"
        if w in base_tokens or alt in base_tokens:
            return True
    return False

# -------------------------
# Replace/insert Summary as bullets
# -------------------------
KNOWN_HEADINGS = [
    "Profile Summary", "Professional Summary", "Summary",
    "Core Skills", "Core Competencies", "Skills", "Technical Skills",
    "Work Experience", "Experience", "Projects", "Education",
    "Certifications", "Achievements", "Publications", "Awards"
]
# Accept headings with optional punctuation (colon/dash)
_HEADING_RE = r"(?im)^\s*(profile\s*summary|professional\s*summary|summary)\s*[:\-–—]?\s*$"
_NEXT_HEADING_RE = r"(?im)^\s*(profile\s*summary|professional\s*summary|summary|core\s*skills|core\s*competencies|skills|technical\s*skills|work\s*experience|experience|education|projects|certifications|awards|publications)\s*[:\-–—]?\s*$"

# Placeholders used to freeze summary position during tailoring
_SUMMARY_START_PH = "\n<<<KEEP_SUMMARY_POSITION_START>>>\n"
_SUMMARY_END_PH   = "\n<<<KEEP_SUMMARY_POSITION_END>>>\n"
def _find_summary_bounds(text: str):
    """
    Return (head_start, head_end, block_end, heading_text) for the first Summary section.
    head_* are the bounds for heading line; block_end is end of the entire section body.
    If not found, return (-1, -1, -1, "").
    """
    if not text:
        return -1, -1, -1, ""
    m = re.search(_HEADING_RE, text)
    if not m:
        return -1, -1, -1, ""
    head_start, head_end = m.start(), m.end()
    after = text[head_end:]
    nxt = re.search(_NEXT_HEADING_RE, after)
    block_end = head_end + (nxt.start() if nxt else len(after))
    heading_text = text[m.start():m.end()].strip()
    return head_start, head_end, block_end, heading_text

def _insert_summary_placeholders(full_text: str):
    """
    Replace ONLY the Summary BODY with placeholders, keep the heading line intact.
    Returns (text_with_placeholders, found: bool)
    """
    hs, he, be, _ = _find_summary_bounds(full_text or "")
    if hs < 0:
        return full_text, False
    head = full_text[:he]
    body = full_text[he:be]
    tail = full_text[be:]
    # Keep the heading; replace the body with placeholders
    return (head.rstrip() + _SUMMARY_START_PH + _SUMMARY_END_PH + tail.lstrip("\n")), True

def _replace_placeholders_with_bullets(full_text: str, bullets_block: str) -> str:
    """
    Replace the placeholder region with the provided bullets (normalized to '• ').
    """
    if not full_text:
        return full_text
    # Normalize bullets
    lines = []
    for ln in (bullets_block or "").splitlines():
        s = ln.strip()
        if not s:
            continue
        if not s.startswith("• "):
            s = "• " + s.lstrip("-").lstrip("•").strip()
        lines.append(s)
    block = "\n".join(lines).strip()
    # Stitch in place
    if _SUMMARY_START_PH in full_text and _SUMMARY_END_PH in full_text:
        return full_text.replace(_SUMMARY_START_PH, "\n").replace(_SUMMARY_END_PH, "\n" + block + "\n", 1).replace(_SUMMARY_END_PH, "")
    return full_text

def _remove_summary_section(text: str) -> str:
    """
    Remove the first Profile/Professional/Summary section (heading + body),
    leaving the rest unchanged. Used so tailor() cannot rewrite the summary.
    """
    t = text or ""
    m = m = re.search(_HEADING_RE, t)
    if not m:
        return t
    start_h = m.start()
    end_h   = m.end()
    after = t[end_h:]
    next_h = re.search(_NEXT_HEADING_RE, after)

    end_block = end_h + (next_h.start() if next_h else len(after))
    return (t[:start_h] + t[end_block:]).strip()

def _extract_existing_summary_block(resume_text: str) -> str:
    """
    Return the raw text inside the first Summary section (Profile/Professional/Summary),
    excluding the heading line and up to the next section heading.
    """
    txt = resume_text or ""
    headings_pattern = r"(?im)^(profile\s*summary|professional\s*summary|summary|core\s*skills|core\s*competencies|skills|technical\s*skills|work\s*experience|experience|education|projects|certifications|awards|publications)\s*$"
    m = re.search(_HEADING_RE, txt)
    if not m:
        return ""
    start = m.end()
    after = txt[start:]
    n = re.search(_NEXT_HEADING_RE, after)
    if n:
        return after[:n.start()].strip()
    return after.strip()

def _extract_core_competencies_block(text: str) -> str:
    """
    Extract the body of the Core Competencies section (excluding heading).
    Handles variations like 'CORE COMPETENCIES:', 'Core-Competencies', etc.
    """
    if not text:
        return ""

    # Match heading in multiple variations
    pattern = re.compile(r"(?im)^\s*core[\s\-_:]*competencies\s*[:\-–—]?\s*$")
    m = pattern.search(text)
    if not m:
        return ""

    after = text[m.end():]

    # Look for the next section heading
    nxt = re.search(
        r"(?im)^\s*(skills|technical\s*skills|work\s*experience|experience|education|projects|certifications|awards|publications|summary)\s*[:\-–—]?\s*$",
        after
    )

    section_text = after[:nxt.start()].strip() if nxt else after.strip()
    return section_text




def _replace_summary_with_bullets(full_text: str, bullets_text: str) -> str:
    """
    Replace the first 'Profile Summary'/'Professional Summary' block with the new bullets.
    If no summary heading exists, insert a 'Profile Summary' + bullets after the first non-empty line.
    """
    if not (full_text and bullets_text and bullets_text.strip()):
        return full_text

    # Normalize bullets to start with "• "
    lines = []
    for ln in bullets_text.splitlines():
        s = ln.strip()
        if not s:
            continue
        if not s.startswith("• "):
            s = "• " + s.lstrip("-").lstrip("•").strip()
        lines.append(s)
    bullet_block = "\n".join(lines)

    # Find summary heading case-insensitively
    m = re.search(_HEADING_RE, full_text or "")
    if m:
        start = m.start()
        end = m.end()
        after = full_text[end:]
        nxt = re.search(_NEXT_HEADING_RE, after)

        section_end = end + (nxt.start() if nxt else len(after))
        head = full_text[:end].rstrip()  # keep the heading line
        tail = full_text[section_end:].lstrip("\n")
        return (head + "\n" + bullet_block + "\n\n" + tail).strip()

    # If there is no summary heading, insert after the first non-empty line
    parts = (full_text or "").splitlines()
    idx = 0
    while idx < len(parts) and not parts[idx].strip():
        idx += 1
    insert_at = min(len(parts), idx + 1)
    new_lines = parts[:insert_at] + ["", "Profile Summary", bullet_block, ""] + parts[insert_at:]
    return "\n".join(new_lines).strip()


# -------------------------
# Streamlit UI
# -------------------------
st.set_page_config(page_title="API-only Resume Tailor (v8 final)", page_icon="🧰", layout="wide")
st.title("🧰 API-only Resume Tailor (v8 final)")
st.caption("v8 layout • Full content read • Colored headings in DOCX/PDF • Plain-text output (no ##/**)")

# -------------------------
# Sidebar
# -------------------------
with st.sidebar:
    st.header("🔑 Provider & Model")
    provider = st.selectbox("Provider", ["openai","gemini","anthropic"], index=1)
    model = st.text_input("Model name (optional)", value="")
    temperature = st.slider("Temperature", 0.0, 1.0, 0.2, 0.05)
    max_tokens = st.slider("Max tokens (output cap)", 256, 8192, 3000, 64)

    st.header("🔐 API Keys")
    openai_key = st.text_input("OpenAI API Key", type="password")
    gemini_key = st.text_input("Gemini API Key", type="password")
    anthropic_key = st.text_input("Anthropic API Key", type="password")
    keys = {"openai": openai_key.strip(), "gemini": gemini_key.strip(), "anthropic": anthropic_key.strip()}

# -------------------------
# Upload or paste helper
# -------------------------
def read_textarea_or_file(label: str, key_text: str, key_file: str) -> str:
    txt = st.session_state.get(key_text, "")
    up = st.file_uploader(label, type=["txt","md","pdf","docx"], key=key_file)
    if up is not None:
        ext = up.name.lower().split(".")[-1]
        if ext in ("txt","md"):
            txt = up.read().decode("utf-8", errors="ignore")
        elif ext == "pdf":
            try:
                from pdfminer.high_level import extract_text
            except Exception:
                # some envs use slightly different import path; fallback above usually works
                from pdfminer.high_level import extract_text
            import tempfile
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                    tmp.write(up.read()); tmp.flush()
                    txt = extract_text(tmp.name) or ""
            except Exception:
                from PyPDF2 import PdfReader
                up.seek(0); reader = PdfReader(up)
                txt = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext == "docx":
            from docx import Document
            d = Document(up)
            paras = [p.text for p in d.paragraphs]
            for tbl in d.tables:
                for row in tbl.rows:
                    paras.append(" | ".join(cell.text for cell in row.cells))
            txt = "\n".join(paras)
    txt = st.text_area(f"Or paste {label.lower()} text here", value=txt, height=240, key=key_text)
    return txt

# -------------------------
# Inputs
# -------------------------
col1, col2 = st.columns(2)
with col1:
    st.subheader("Resume")
    resume_text = read_textarea_or_file("Resume", "resume_text", "resume_file")
with col2:
    st.subheader("Job Description")
    jd_text = read_textarea_or_file("Job Description", "jd_text", "jd_file")

# Always use text boxes as source of truth
resume_text = st.session_state.get("resume_text", "") or resume_text
jd_text = st.session_state.get("jd_text", "") or jd_text

# -------------------------
# Main
# -------------------------
if resume_text.strip() and jd_text.strip():
    st.divider()
    st.subheader("Analysis")

    report = analyze(resume_text, jd_text)

    if st.button("Enhance contacts with AI", key="btn_contacts_ai"):
        try:
            st.session_state["ai_contacts"] = extract_contacts_llm(
                resume_text, provider_pref=provider, model_name=(model or None), keys=keys
            )
            st.success("Contacts enhanced.")
        except Exception as e:
            st.error(str(e))

    ai_contacts = st.session_state.get("ai_contacts", {})
    merged = {**report["contacts"], **{k: v for k, v in ai_contacts.items() if v}}

    st.markdown("**Contacts**")
    c1, c2, c3, c4, c5 = st.columns(5)
    with c1: merged["name"] = st.text_input("Name", merged.get("name",""))
    with c2: merged["email"] = st.text_input("Email", merged.get("email",""))
    with c3: merged["phone"] = st.text_input("Phone", merged.get("phone",""))
    with c4: merged["linkedin"] = st.text_input("LinkedIn", merged.get("linkedin",""))
    with c5: merged["github"] = st.text_input("GitHub", merged.get("github",""))
    st.session_state["override_contacts"] = merged

    cA, cB, cC = st.columns([1,1,2])
    with cA: st.metric("Match Score", f"{report['match']['match_score']}%")
    with cB: st.metric("Readability (FRE)", report["readability"]["flesch_reading_ease"])
    with cC:
        st.write("ATS warnings:")
        for w in report["ats"]["warnings"]:
            st.write(f"• {w}")

    # -----------------
    # LLM Keyword Optimizer
    # -----------------
    st.subheader("LLM Keyword Optimizer")
    if st.button("Extract ranked keywords with AI", key="btn_kw_extract"):
        try:
            kw = extract_keywords_llm(
                resume_text, jd_text,
                provider_pref=provider, model_name=(model or None),
                temperature=temperature, max_tokens=min(max_tokens, 1200), keys=keys
            )
            st.session_state["kw_llm"] = kw
        except Exception as e:
            st.error(str(e))

    kw_obj = st.session_state.get("kw_llm")

    if kw_obj and "_raw_extraction" in kw_obj:
        with st.expander("🔍 Raw Extraction from LLM"):
            st.text(kw_obj["_raw_extraction"])




    def _fallback_missing_and_weak(kw_obj, resume_text, jd_text):
        """
        Extract true skill gaps: tools/technologies from JD not already in Top Keywords.
        No verbs or filler words.
        """
        gap_terms = []
        try:
            # Capture capitalized tech terms / acronyms from JD
            jd_tokens = re.findall(r"\b[A-Z][A-Za-z0-9+\-_/]{2,}\b", jd_text)
            jd_tokens = [t for t in jd_tokens if len(t) > 2]

            # Already covered keywords
            seen_keywords = { (item.get("term") or "").lower() for item in (kw_obj.get("keywords") or []) }

            # Filter out those already present
            gap_terms = [t for t in jd_tokens if t.lower() not in seen_keywords]

            # Deduplicate and sort
            gap_terms = sorted(set(gap_terms))
        except Exception:
            gap_terms = []

        # Weak terms: keep existing logic (or leave empty if you don’t want it at all)
        weak_terms = []
        return gap_terms, weak_terms


    # Build target keywords from Optimizer (ranked + gaps), deduped, excluding ones already in resume
    def build_target_keywords_from_optimizer(kw_obj, resume_text: str, jd_text: str):
        if not kw_obj:
            return []

        token_re = re.compile(r"[A-Za-z0-9#+.]+")

        def _canon_local(s: str) -> str:
            return " ".join(t.lower() for t in token_re.findall(s or ""))

        def _tok_seq_local(s: str):
            return [t.lower() for t in token_re.findall(s or "")]

        resume_tokens = _tok_seq_local(resume_text)
        resume_compact = "".join(resume_tokens)

        def _present(term: str) -> bool:
            kt = _tok_seq_local(term)
            if not kt:
                return False
            L = len(kt)
            for i in range(0, len(resume_tokens) - L + 1):
                if resume_tokens[i:i+L] == kt:
                    return True
            k_comp = "".join(kt)
            if k_comp and k_comp in resume_compact:
                return True
            if L == 1 and len(kt[0]) > 3:
                base = kt[0]
                alt = base[:-1] if base.endswith("s") else base + "s"
                if base in resume_tokens or alt in resume_tokens:
                    return True
            return False

        outlets_ordered = []
        for item in (kw_obj.get("keywords") or []):
            base = (item.get("term") or "").strip()
            if not base:
                continue
            variants = [v.strip() for v in (item.get("variants") or []) if v and v.strip()]
            outlets_ordered.append(base)
            outlets_ordered.extend(variants)

        gaps_terms = list(kw_obj.get("missing") or [])
        for g in gaps_terms:
            g = (g or "").strip()
            if g:
                outlets_ordered.append(g)

        seen, deduped = set(), []
        for t in outlets_ordered:
            c = _canon_local(t)
            if c and c not in seen:
                seen.add(c)
                deduped.append(t)

        target = [t for t in deduped if not _present(t)]
        return target

    if kw_obj:
        st.write(kw_obj.get("summary",""))
        colk1, colk2 = st.columns(2)
        with colk1:
            st.markdown("**Top Keywords (ranked)**")

            # make sure kw_obj is dict and has 'keywords'
            keywords = []
            if isinstance(kw_obj, dict):
                keywords = kw_obj.get("keywords", [])

            if keywords and isinstance(keywords, list):
                for item in keywords:
                    term = item.get("term", "").strip()
                    if not term:
                        continue
                    cat = item.get("category", "general")
                    variants = item.get("variants", [])
                    rank = item.get("rank", "?")
                    suffix = f" · variants: {', '.join(variants)}" if variants else ""
                    st.write(f"{rank}. **{term}** · _{cat}_{suffix}")
            else:
                st.error("⚠️ No parsed keywords available from LLM.")
                st.text("=== RAW JSON FROM LLM ===\n" + str(kw_obj.get("_raw_json", "")))


        with colk2:
            st.markdown("**Gaps**")
            try:
                gaps = extract_gaps(resume_text, kw_obj)
            except Exception:
                gaps = []
            st.write("Gaps:", ", ".join(gaps) if gaps else "—")
            st.caption("🔍 These are Top Keywords missing from your resume.")



    # -----------------
    # Keyword Sentence Generator (ATS-friendly) — SEPARATE EDITOR
    # -----------------
    st.subheader("Keyword Sentence Generator (ATS-friendly)")
    st.caption("Generates concise 'Core Competencies' bullets using Top Keywords (ranked) + Gaps. Edit here and Save; Tailor will blend them in.")

    col_gen, col_clear = st.columns([1,1])
    with col_gen:
        if st.button("Generate ATS-friendly keyword sentences", key="btn_kw_sentences_generate"):
            if not kw_obj:
                st.warning("Please run the LLM Keyword Optimizer first.")
            else:
                # Build a full list and classify present/new (variants-aware)
                TOKEN_RE = re.compile(r"[A-Za-z0-9#+.]+")
                def _tok_seq_local2(s: str):
                    return [t.lower() for t in TOKEN_RE.findall(s or "")]
                def _canon_local2(s: str) -> str:
                    return " ".join(_tok_seq_local2(s))

                ranked_terms = []
                for it in (kw_obj.get("keywords") or []):
                    term = (it.get("term") or "").strip()
                    if term:
                        ranked_terms.append(term)

                gaps_terms = list(kw_obj.get("missing") or [])
                if not gaps_terms:
                    try:
                        fallback_missing, _fw = _fallback_missing_and_weak(kw_obj, resume_text, jd_text)
                    except Exception:
                        fallback_missing = []
                    gaps_terms = list(kw_obj.get("missing") or [])

                seen, target_all = set(), []
                for t in ranked_terms + gaps_terms:
                    c = _canon_local2(t)
                    if c and c not in seen:
                        seen.add(c)
                        target_all.append(t)

                variants_map = {}
                for it in (kw_obj.get("keywords") or []):
                    term = (it.get("term") or "").strip()
                    if not term:
                        continue
                    forms = [term] + [v for v in (it.get("variants") or []) if v and v.strip()]
                    expanded = set()
                    for f in forms:
                        f = f.strip()
                        if not f:
                            continue
                        expanded.add(f)
                        expanded.add(f.replace("-", " "))
                        expanded.add(f.replace("/", " "))
                    variants_map[_canon_local2(term)] = list(expanded)

                presence_text = (st.session_state.get("tailored_edit") or resume_text or "")
                text_tokens = _tok_seq_local2(presence_text)

                def _has_seq(tokens, cand: str) -> bool:
                    seq = _tok_seq_local2(cand)
                    L = len(seq)
                    if L == 0:
                        return False
                    for i in range(0, len(tokens) - L + 1):
                        if tokens[i:i+L] == seq:
                            return True
                    if L == 1 and len(seq[0]) > 3:
                        w = seq[0]
                        alt = w[:-1] if w.endswith("s") else w + "s"
                        return (w in tokens) or (alt in tokens)
                    return False

                present_keywords, new_keywords = [], []
                for t in target_all:
                    key = _canon_local2(t)
                    cand_forms = variants_map.get(key, [t, t.replace("-", " "), t.replace("/", " ")])
                    seen_forms, cand_forms_dedup = set(), []
                    for f in cand_forms:
                        f = (f or "").strip()
                        if not f:
                            continue
                        kf = f.lower()
                        if kf in seen_forms:
                            continue
                        seen_forms.add(kf)
                        cand_forms_dedup.append(f)
                    found = any(_has_seq(text_tokens, f) for f in cand_forms_dedup)
                    (present_keywords if found else new_keywords).append(t)

                with st.expander("🔍 New keywords for your resume (not currently found)", expanded=True):
                    st.write(", ".join(new_keywords) if new_keywords else "— none —")
                with st.expander("✅ Already present (will NOT be generated again)", expanded=False):
                    st.write(", ".join(present_keywords) if present_keywords else "— none —")

                target_for_sentences = new_keywords

                if not target_for_sentences:
                    st.info("All optimizer keywords already appear in the resume. Nothing to add.")
                else:
                    bullets_raw = generate_keyword_sentences(
                        resume_text=presence_text,
                        jd_text=jd_text,
                        target_keywords=target_for_sentences,
                        provider_pref=provider,
                        model_name=(model or None),
                        temperature=temperature,
                        max_tokens=min(max_tokens, 900),
                        keys=keys
                    )

                    try:
                        polished = polish_keyword_sentences(
                            resume_text=presence_text,
                            bullets_text=(bullets_raw or "").strip(),
                            jd_text=jd_text,
                            provider_pref=provider,
                            model_name=(model or None),
                            temperature=temperature,
                            max_tokens=min(max_tokens, 800),
                            keys=keys
                        )
                        out_text = polished
                    except Exception:
                        out_text = (bullets_raw or "")

                    TOKEN_RE2 = re.compile(r"[A-Za-z0-9#+.]+")
                    def _canon_line(s: str) -> str:
                        return " ".join(t.lower() for t in TOKEN_RE2.findall(s or ""))

                    seen_lines, lines_out = set(), []
                    for ln in (out_text.splitlines() if out_text else []):
                        s = ln.strip()
                        if not s:
                            continue
                        c = _canon_line(s)
                        if c in seen_lines:
                            continue
                        seen_lines.add(c)
                        if not s.startswith("• "):
                            s = "• " + s.lstrip("-").lstrip("•").strip()
                        lines_out.append(s)

                    st.session_state["kw_sentences_edit"] = "\n".join(lines_out).strip()
                    st.success("Generated ATS-friendly keyword sentences. Review below, edit, then click Save.")

    with col_clear:
        if st.button("Clear keyword sentences", key="btn_kw_sentences_clear"):
            st.session_state["kw_sentences_edit"] = ""
            st.session_state["kw_sentences_saved_text"] = ""
            st.info("Keyword sentences cleared.")

    kw_edit = st.text_area("Keyword Sentences (editable, plain text)", key="kw_sentences_edit", height=220)

    if st.button("💾 Save keyword sentences", key="btn_kw_sentences_save"):
        st.session_state["kw_sentences_saved_text"] = (kw_edit or "").strip()
        st.success("Saved. Tailor with LLM will integrate these into the resume.")

    # -----------------
    # Tailor with LLM (API-only) — integrates ONLY Resume + SAVED Keyword Sentences
    # -----------------
    st.divider()
    st.subheader("Tailor with LLM (API-only)")
    st.caption("Integrates your Resume + SAVED Keyword Sentence Generator text so it reads like original experience (no copy-paste feel), avoids duplicates, and refines existing mentions.")

    if st.button("Generate tailored resume", type="primary", key="btn_tailor_generate"):
        if not resume_text.strip():
            st.warning("Please paste or upload your resume text first.")
        else:
            try:
                # 1) Read SAVED keyword sentences
                saved_kw_sentences = (st.session_state.get("kw_sentences_saved_text", "") or "").strip()

                # 2) Split & clean
                raw_lines = [ln.strip() for ln in (saved_kw_sentences.splitlines() if saved_kw_sentences else [])]
                raw_lines = [ln for ln in raw_lines if ln]

                # 3) Decide: add vs refine (NO placement logic here)
                lines_to_add, refine_hints, seen_lines = [], [], set()
                for ln in raw_lines:
                    ln_clean = ln.lstrip("•").lstrip("-").strip()
                    if not ln_clean:
                        continue
                    c = _canon(ln_clean)
                    if c in seen_lines:
                        continue
                    seen_lines.add(c)

                    if _present_line(resume_text, ln_clean):
                        refine_hints.append(ln_clean)
                    else:
                        if not ln_clean.startswith("• "):
                            ln_clean = "• " + ln_clean
                        lines_to_add.append(ln_clean)

                # 4) Build input for LLM: base resume + integration notes (AI places content)
                blocks = [resume_text.strip()]
                if refine_hints or lines_to_add:
                    guidance = [
                        "",
                        "Integration Notes (for model):",
                        "- Integrate the following without duplicating existing content.",
                        "- If a theme already exists, refine in-place; do not add a new line.",
                        "- Choose the most appropriate existing section (e.g., Core Competencies/Skills, Technical Skills, or a relevant role).",
                        "- NEVER place added lines at the very top of the document or the very end.",
                        "- Treat ALL added lines below as 'Core Competencies' content; do NOT place them under Work Experience or Projects.",
                    ]
                    if refine_hints:
                        guidance += ["", "Refine these existing themes:"]
                        guidance += [f"- {h}" for h in refine_hints]
                    if lines_to_add:
                        guidance += ["", "Add these lines naturally (responsibility-style):"]
                        guidance += lines_to_add
                    blocks.append("\n".join(guidance))

                base_resume_for_llm = "\n\n".join(blocks).strip()
                
                # --- Freeze summary position with placeholders (keep heading intact) ---
                orig_summary_block = _extract_existing_summary_block(resume_text)
                resume_frozen, has_summary = _insert_summary_placeholders(base_resume_for_llm)

                # Strengthen the instruction to the model in the integration notes
                if "Integration Notes (for model):" in resume_frozen:
                    resume_frozen += (
                        "\n- DO NOT move or delete the markers '<<<KEEP_SUMMARY_POSITION_START>>>' "
                        "and '<<<KEEP_SUMMARY_POSITION_END>>>'."
                        "\n- Keep the exact section order as provided."
                    )

                override_contacts = st.session_state.get("override_contacts")

                # Tailor the resume WITHOUT letting the model touch the frozen summary body
                tailored = tailor(
                    resume_frozen,
                    jd_text,
                    provider_preference=provider,
                    model_name=(model or None),
                    temperature=temperature,
                    max_tokens=max_tokens,
                    keys=keys,
                    target_keywords=[],
                    override_contacts=override_contacts
                )

                final_txt = sanitize_markdown(tailored)

                
                # --- Role-aligned bullet Summary via LLM ---
                if has_summary and orig_summary_block.strip():
                    bullets_text = generate_summary_bullets(
                        resume_text=orig_summary_block,   # only the original summary text
                        jd_text=jd_text,                  # align wording to the role
                        focus="summary",
                        provider_pref=provider,
                        model_name=(model or None),
                        temperature=temperature,
                        max_tokens=min(max_tokens, 1000),  # allow longer bullets
                        keys=keys,
                    )
                    if bullets_text:
                        final_txt = _replace_placeholders_with_bullets(final_txt, bullets_text)

                # Remove any leftover placeholders just in case
                final_txt = final_txt.replace(_SUMMARY_START_PH, "\n").replace(_SUMMARY_END_PH, "\n")

                # --- Strict Core Competencies (merge + polish) ---
                saved_kw_sentences = (st.session_state.get("kw_sentences_saved_text", "") or "").strip()
                orig_core_block = _extract_core_competencies_block(resume_text)
                st.write("🔍 Extracted Core Competencies from resume:", orig_core_block)   # <--- ADD HERE

                if orig_core_block or saved_kw_sentences:
                    from pipeline import polish_core_competencies, replace_core_competencies
                    polished_core = polish_core_competencies(
                        original_bullets=orig_core_block,
                        new_bullets=saved_kw_sentences,
                        provider_pref=provider,
                        model_name=(model or None),
                        temperature=temperature,
                        max_tokens=min(max_tokens, 900),
                        keys=keys,
                    )
                    if polished_core:
                        final_txt = replace_core_competencies(final_txt, polished_core)




                # Persist to editor
                st.session_state["tailored_text"] = final_txt
                st.session_state["tailored_edit"]  = final_txt
                st.session_state["tailored_saved"] = False

                if lines_to_add and refine_hints:
                    st.success("Tailored resume generated. New keyword sentences were integrated naturally and existing mentions were refined. Review and click Save before exporting.")
                elif lines_to_add:
                    st.success("Tailored resume generated. Your keyword sentences were integrated without duplicates. Review and click Save before exporting.")
                elif refine_hints:
                    st.success("Tailored resume generated. Existing keyword mentions were refined (no duplicates added). Review and click Save before exporting.")
                else:
                    st.info("No new keyword sentences found and nothing specific to refine. The resume was still tailored for structure and clarity.")
            except Exception as e:
                st.error(str(e))

    # -------------------------
    # Editor (always visible) + Save + Downloads
    # -------------------------
    edited_text = st.text_area("Tailored resume (plain text)", key="tailored_edit", height=420)

    if st.button("💾 Save", key="btn_tailor_save"):
        def _normalize_save(s: str) -> str:
            s = re.sub(r"[ \t]+$", "", s, flags=re.MULTILINE)
            s = re.sub(r"\n{3,}", "\n\n", s)
            return s.strip()
        st.session_state["tailored_text"] = _normalize_save(edited_text or "")
        st.session_state["tailored_saved"] = True
        st.success("Saved. Exports will use your edited text.")

    saved_text = (st.session_state.get("tailored_text", "") or "").strip()
    colx1, colx2 = st.columns(2)
    with colx1:
        if st.session_state.get("tailored_saved") and saved_text:
            out_path = export_docx(saved_text, "tailored_resume.docx")
            with open(out_path, "rb") as f:
                st.download_button("⬇️ Download DOCX", f, file_name="tailored_resume.docx", key="dl_docx")
        else:
            st.info("Click Save to enable DOCX download.")
    with colx2:
        if st.session_state.get("tailored_saved") and saved_text:
            out_path = export_pdf(saved_text, "tailored_resume.pdf")
            with open(out_path, "rb") as f:
                st.download_button("⬇️ Download PDF", f, file_name="tailored_resume.pdf", key="dl_pdf")
        else:
            st.info("Click Save to enable PDF download.")

    # -----------------
    # ATS Scan (AI) + local reconciliation
    # -----------------
    st.divider()
    st.subheader("ATS Scan (Keyword Coverage vs Final Resume)")

    if st.button("Run ATS analysis on edited resume", key="btn_ats_run"):
        final_text = (st.session_state.get("tailored_edit", "") or "").strip()
        kw_obj_now = st.session_state.get("kw_llm") or {}

        if not final_text:
            st.warning("Please add or generate resume content first.")
        elif not kw_obj_now:
            st.warning("Please run the LLM Keyword Optimizer first.")
        else:
            ats_llm = extract_ats_llm_from_optimizer(
                resume_text=final_text,
                optimizer_obj=kw_obj_now,
                provider_pref=provider,
                model_name=(model or None),
                temperature=temperature,
                max_tokens=max_tokens,
                keys=keys,
                jd_text=jd_text
            )
            st.session_state["final_ats_llm"] = ats_llm

            # Deterministic reconciliation (token/variant aware + compact match)
            token_re = re.compile(r"[A-Za-z0-9#+.]+")
            def _canon_local3(s: str) -> str:
                return " ".join(t.lower() for t in token_re.findall(s or ""))

            ranked = [(it.get("term") or "").strip() for it in (kw_obj_now.get("keywords") or []) if (it.get("term") or "").strip()]
            gaps = list(kw_obj_now.get("missing") or [])

            seen_terms, ordered_terms = set(), []
            for t in (ranked + gaps):
                c = _canon_local3(t)
                if c and c not in seen_terms:
                    seen_terms.add(c); ordered_terms.append(t)

            variants_map = {}
            for it in (kw_obj_now.get("keywords") or []):
                term = (it.get("term") or "").strip()
                if term:
                    variants_map[_canon_local3(term)] = [v.strip() for v in (it.get("variants") or []) if v and v.strip()]

            final_tokens = [t.lower() for t in token_re.findall(final_text)]
            final_compact = "".join(final_tokens)

            def _tok_seq3(s: str):
                return [t.lower() for t in token_re.findall(s or "")]

            present, missing, coverage = [], [], []
            for term in ordered_terms:
                cand_list = [term] + (variants_map.get(_canon_local3(term), []))
                found = False
                found_pos = -1
                found_len = 0

                for cand in cand_list:
                    tt = _tok_seq3(cand)
                    if not tt:
                        continue

                    L = len(tt)
                    for i in range(0, len(final_tokens) - L + 1):
                        if final_tokens[i:i+L] == tt:
                            found, found_pos, found_len = True, i, L
                            break
                    if found:
                        break

                    t_comp = "".join(tt)
                    if t_comp and t_comp in final_compact:
                        found, found_pos, found_len = True, -1, L
                        break

                    if L == 1 and len(tt[0]) > 3:
                        base = tt[0]
                        alt = base[:-1] if base.endswith("s") else base + "s"
                        if base in final_tokens or alt in final_tokens:
                            found, found_pos, found_len = True, -1, 1
                            break

                if found:
                    if found_pos >= 0:
                        spans = [(m.group(0), m.start(), m.end()) for m in re.finditer(r"[A-Za-z0-9#+.]+", final_text)]
                        s = spans[found_pos][1] if 0 <= found_pos < len(spans) else 0
                        e_idx = min(found_pos + max(1, found_len) - 1, len(spans) - 1)
                        e = spans[e_idx][2] if spans else s
                        s = max(0, s - 20); e = min(len(final_text), e + 20)
                        ev = final_text[s:e].replace("\n", " ").strip()
                    else:
                        ev = ""
                    present.append(term)
                    coverage.append({"term": term, "present": True, "evidence": ev})
                else:
                    missing.append(term)
                    coverage.append({"term": term, "present": False, "evidence": ""})

            total = max(1, len(ordered_terms))
            score = round(100 * len(present) / total)
            st.session_state["final_ats_llm"] = {
                **(st.session_state.get("final_ats_llm") or {}),
                "score": score,
                "present": present,
                "missing": missing,
                "coverage": coverage,
            }

    # Display ATS results
    ats_llm = st.session_state.get("final_ats_llm")
    if ats_llm:
        c1, c2 = st.columns([1, 2])
        with c1:
            st.metric("AI ATS Keyword Score", f"{ats_llm.get('score',0)}%")
        with c2:
            st.write("Suggestions (where & how to add missing keywords):")
            sugg = ats_llm.get("suggestions") or []
            miss = set(ats_llm.get("missing") or [])
            shown = False
            for s in sugg:
                term = s.get("term","")
                if term in miss:
                    shown = True
                    st.write(f"• {term} → {s.get('section','Core Competencies')}: {s.get('how','')}")
            if not shown:
                st.write("—")

        colp, colm = st.columns(2)
        with colp:
            st.markdown("**Present keywords**")
            pres = ats_llm.get("present") or []
            st.write(", ".join(pres) if pres else "—")
        with colm:
            st.markdown("**Missing keywords**")
            miss = ats_llm.get("missing") or []
            st.write(", ".join(miss) if miss else "—")

        cov = ats_llm.get("coverage") or []
        if cov:
            st.markdown("**Coverage details (sample)**")
            for row in cov[:12]:
                t = row.get("term",""); p = "✅" if row.get("present") else "❌"
                ev = row.get("evidence","")
                st.write(f"{p} {t}: {ev}")

else:
    st.info("Upload or paste both the Resume and the Job Description to begin.")
