from io import BytesIO
from typing import Dict, Any, List

# -------- DOCX ----------
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

def _add_bullets(doc, bullets: List[str]):
    for b in bullets:
        if b and b.strip():
            p = doc.add_paragraph(b.strip(), style="List Bullet")

def build_docx(resume: Dict[str, Any]) -> bytes:
    doc = Document()

    # Base styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Header
    name = resume.get("name") or "Your Name"
    head = doc.add_paragraph()
    run = head.add_run(name)
    run.bold = True
    run.font.size = Pt(18)
    head.alignment = WD_ALIGN_PARAGRAPH.LEFT

    c = resume.get("contact", {})
    contact_line = " • ".join(filter(None, [
        c.get("email"), c.get("phone"), c.get("location"),
        " | ".join(c.get("links", []) or [])
    ]))
    if contact_line:
        p = doc.add_paragraph(contact_line)
        p.style = doc.styles['Normal']

    # Summary
    if resume.get("summary"):
        _add_heading(doc, "Professional Summary", level=2)
        doc.add_paragraph(resume["summary"])

    # Skills
    skills = resume.get("skills") or []
    if skills:
        _add_heading(doc, "Skills", level=2)
        doc.add_paragraph(", ".join(skills))

    # Experience
    exp = resume.get("experience") or []
    if exp:
        _add_heading(doc, "Experience", level=2)
        for r in exp:
            title = " • ".join(filter(None, [r.get("title"), r.get("company"), r.get("location")]))
            dates = " - ".join(filter(None, [r.get("start"), r.get("end")]))
            doc.add_paragraph(title).runs[0].bold = True
            if dates:
                doc.add_paragraph(dates)
            _add_bullets(doc, r.get("bullets") or [])

    # Projects
    projs = resume.get("projects") or []
    if projs:
        _add_heading(doc, "Projects", level=2)
        for p in projs:
            hdr = " • ".join(filter(None, [p.get("name"), p.get("role")]))
            doc.add_paragraph(hdr).runs[0].bold = True
            _add_bullets(doc, p.get("bullets") or [])

    # Education
    edu = resume.get("education") or []
    if edu:
        _add_heading(doc, "Education", level=2)
        for e in edu:
            hdr = " • ".join(filter(None, [e.get("degree"), e.get("school"), e.get("year")]))
            doc.add_paragraph(hdr).runs[0].bold = True
            _add_bullets(doc, e.get("details") or [])

    # Certifications
    certs = resume.get("certifications") or []
    if certs:
        _add_heading(doc, "Certifications", level=2)
        _add_bullets(doc, certs)

    # Extras
    extras = resume.get("extras") or {}
    for section, items in extras.items():
        if items:
            _add_heading(doc, section.capitalize(), level=2)
            _add_bullets(doc, items)

    # Export
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

# -------- PDF ----------
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def build_pdf(resume: Dict[str, Any]) -> bytes:
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Section", fontSize=12, leading=14, spaceBefore=10, spaceAfter=6, textColor="#000000", underlineWidth=0.5))
    styles.add(ParagraphStyle(name="Header", fontSize=18, leading=22, spaceAfter=6, bold=True))
    styles.add(ParagraphStyle(name="Subtle", fontSize=9, leading=11, textColor="#333333"))

    flow = []

    name = resume.get("name") or "Your Name"
    flow.append(Paragraph(name, styles["Heading1"]))
    c = resume.get("contact", {})
    contact_line = " • ".join(filter(None, [
        c.get("email"), c.get("phone"), c.get("location"),
        " | ".join(c.get("links", []) or [])
    ]))
    if contact_line:
        flow.append(Paragraph(contact_line, styles["Subtle"]))
    flow.append(Spacer(1, 8))

    if resume.get("summary"):
        flow.append(Paragraph("Professional Summary", styles["Section"]))
        flow.append(Paragraph(resume["summary"], styles["BodyText"]))

    def add_list(title, items):
        if not items: return
        flow.append(Spacer(1, 6))
        flow.append(Paragraph(title, styles["Section"]))
        lst = ListFlowable([ListItem(Paragraph(i, styles["BodyText"])) for i in items], bulletType="bullet")
        flow.append(lst)

    skills = resume.get("skills") or []
    if skills:
        add_list("Skills", [", ".join(skills)])

    exp = resume.get("experience") or []
    if exp:
        flow.append(Spacer(1, 6))
        flow.append(Paragraph("Experience", styles["Section"]))
        for r in exp:
            hdr = " • ".join(filter(None, [r.get("title"), r.get("company"), r.get("location")]))
            dates = " - ".join(filter(None, [r.get("start"), r.get("end")]))
            flow.append(Paragraph(f"<b>{hdr}</b>", styles["BodyText"]))
            if dates:
                flow.append(Paragraph(dates, styles["Subtle"]))
            bullets = r.get("bullets") or []
            if bullets:
                lst = ListFlowable([ListItem(Paragraph(b, styles["BodyText"])) for b in bullets], bulletType="bullet")
                flow.append(lst)

    projs = resume.get("projects") or []
    if projs:
        flow.append(Spacer(1, 6))
        flow.append(Paragraph("Projects", styles["Section"]))
        for p in projs:
            hdr = " • ".join(filter(None, [p.get("name"), p.get("role")]))
            flow.append(Paragraph(f"<b>{hdr}</b>", styles["BodyText"]))
            bullets = p.get("bullets") or []
            if bullets:
                lst = ListFlowable([ListItem(Paragraph(b, styles["BodyText"])) for b in bullets], bulletType="bullet")
                flow.append(lst)

    edu = resume.get("education") or []
    if edu:
        flow.append(Spacer(1, 6))
        flow.append(Paragraph("Education", styles["Section"]))
        for e in edu:
            hdr = " • ".join(filter(None, [e.get("degree"), e.get("school"), e.get("year")]))
            flow.append(Paragraph(f"<b>{hdr}</b>", styles["BodyText"]))
            bullets = e.get("details") or []
            if bullets:
                lst = ListFlowable([ListItem(Paragraph(b, styles["BodyText"])) for b in bullets], bulletType="bullet")
                flow.append(lst)

    certs = resume.get("certifications") or []
    if certs:
        add_list("Certifications", certs)

    extras = resume.get("extras") or {}
    for section, items in extras.items():
        add_list(section.capitalize(), items)

    doc.build(flow)
    return buf.getvalue()
